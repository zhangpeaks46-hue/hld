import os
import docx
from PyPDF2 import PdfReader
from typing import List, Dict, Any

class DocumentProcessor:
    def __init__(self, file_path: str, file_type: str):
        self.file_path = file_path
        self.file_type = file_type
        self.content = ""
        self.metadata = {}
    
    def extract_text(self) -> str:
        if self.file_type == ".docx":
            return self._extract_docx_text()
        elif self.file_type == ".pdf":
            return self._extract_pdf_text()
        return ""
    
    def _extract_docx_text(self) -> str:
        try:
            doc = docx.Document(self.file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return "\n".join(full_text)
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def _extract_pdf_text(self) -> str:
        try:
            reader = PdfReader(self.file_path)
            full_text = []
            for page in reader.pages:
                full_text.append(page.extract_text())
            return "\n".join(full_text)
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def analyze_format(self) -> List[Dict[str, Any]]:
        issues = []
        
        if self.file_type == ".docx":
            issues.extend(self._analyze_docx_format())
        elif self.file_type == ".pdf":
            issues.extend(self._analyze_pdf_format())
        
        return issues
    
    def _analyze_docx_format(self) -> List[Dict[str, Any]]:
        issues = []
        try:
            doc = docx.Document(self.file_path)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name not in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                    issues.append({
                        "page": None,
                        "line": i + 1,
                        "issue_type": "格式问题",
                        "description": f"段落使用非标准样式: {paragraph.style.name}",
                        "suggestion": "建议使用标准样式（标题1-3或正文）",
                        "status": "需手动修改"
                    })
            
            for table in doc.tables:
                if len(table.rows) > 50 or len(table.columns) > 10:
                    issues.append({
                        "page": None,
                        "line": None,
                        "issue_type": "表格问题",
                        "description": "表格过大，可能影响阅读体验",
                        "suggestion": "考虑拆分表格或调整布局",
                        "status": "需手动修改"
                    })
            
            if len(issues) == 0:
                issues.append({
                    "page": 1,
                    "line": 1,
                    "issue_type": "格式检查",
                    "description": "文档格式检查完成，未发现明显问题",
                    "suggestion": "文档格式良好",
                    "status": "已通过"
                })
                
        except Exception as e:
            issues.append({
                "page": None,
                "line": None,
                "issue_type": "读取错误",
                "description": f"无法读取文档: {str(e)}",
                "suggestion": "请检查文档是否损坏",
                "status": "错误"
            })
        
        return issues
    
    def _analyze_pdf_format(self) -> List[Dict[str, Any]]:
        issues = []
        try:
            reader = PdfReader(self.file_path)
            num_pages = len(reader.pages)
            
            if num_pages > 100:
                issues.append({
                    "page": None,
                    "line": None,
                    "issue_type": "文档过长",
                    "description": f"文档共 {num_pages} 页，建议拆分",
                    "suggestion": "考虑将长文档拆分为多个文件",
                    "status": "建议"
                })
            
            issues.append({
                "page": 1,
                "line": 1,
                "issue_type": "格式检查",
                "description": f"PDF文档共 {num_pages} 页，已完成格式检查",
                "suggestion": "PDF格式检查完成",
                "status": "已通过"
            })
            
        except Exception as e:
            issues.append({
                "page": None,
                "line": None,
                "issue_type": "读取错误",
                "description": f"无法读取PDF: {str(e)}",
                "suggestion": "请检查PDF是否损坏",
                "status": "错误"
            })
        
        return issues
    
    def proofread_text(self) -> List[Dict[str, Any]]:
        issues = []
        text = self.extract_text()
        
        common_errors = {
            "的": ["地", "得"],
            "地": ["的", "得"],
            "得": ["的", "地"],
            "在": ["再"],
            "再": ["在"],
            "做": ["作"],
            "作": ["做"],
        }
        
        lines = text.split('\n')
        for line_num, line in enumerate(lines):
            for char, similar_chars in common_errors.items():
                if char in line:
                    for similar_char in similar_chars:
                        if similar_char in line:
                            issues.append({
                                "page": None,
                                "line": line_num + 1,
                                "issue_type": "用词混淆",
                                "description": f"发现 '{char}' 和 '{similar_char}' 同时出现，可能存在用词不当",
                                "suggestion": f"请确认使用 '{char}' 还是 '{similar_char}'",
                                "status": "需手动修改"
                            })
        
        if len(text) < 10:
            issues.append({
                "page": None,
                "line": 1,
                "issue_type": "内容过短",
                "description": "文档内容过短，建议补充更多内容",
                "suggestion": "请添加更多内容",
                "status": "建议"
            })
        
        if len(issues) == 0:
            issues.append({
                "page": 1,
                "line": 1,
                "issue_type": "校对完成",
                "description": "文字校对完成，未发现明显问题",
                "suggestion": "文档校对通过",
                "status": "已通过"
            })
        
        return issues
    
    def process_text(self) -> Dict[str, Any]:
        text = self.extract_text()
        
        processed_text = text
        improvements = []
        
        if len(text) > 0:
            improvements.append({
                "type": "字数统计",
                "original_length": len(text),
                "description": "文档处理完成"
            })
            
            if len(text) > 5000:
                improvements.append({
                    "type": "内容优化建议",
                    "description": "文档较长，建议分段处理",
                    "suggestion": "考虑将长文档分成多个章节"
                })
        
        return {
            "processed_text": processed_text,
            "improvements": improvements,
            "original_length": len(text),
            "processed_length": len(processed_text)
        }